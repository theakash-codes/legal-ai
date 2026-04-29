# api/views.py
from ml.predict import predict_text
import io, re, base64, textwrap, traceback, logging, os, json
from django.http import HttpResponse
from django.conf import settings
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from PyPDF2 import PdfReader
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.permissions import AllowAny
from rest_framework import status, generics
from .models import Document, History, GeneratedClause
from .serializers import DocumentSerializer, HistorySerializer
from . import nlp_engine
from . import legal_templates
from . import lawyer_data


logger = logging.getLogger(__name__)

# Register Kannada font
font_path = os.path.join(settings.BASE_DIR, 'static', 'fonts', 'NotoSansKannada-Regular.ttf')
try:
    pdfmetrics.registerFont(TTFont('NotoSansKannada', font_path))
except Exception:
    pass


def sanitize_text(text):
    """Remove XML tags, binary junk, encoded metadata, and non-readable content."""
    if not text:
        return ""
    # Strip XML/HTML tags
    text = re.sub(r'<[^>]+>', ' ', text)
    # Remove XML declarations and processing instructions
    text = re.sub(r'<\?[^?]*\?>', '', text)
    # Remove common DOCX internal paths/references
    text = re.sub(r'word/(?:document|styles|settings|fontTable|numbering|footnotes|endnotes|header\d*|footer\d*|theme/theme\d*|stylesWithEffects)\.xml', '', text)
    text = re.sub(r'word/_rels/[\w.]+', '', text)
    text = re.sub(r'_rels/\.rels', '', text)
    text = re.sub(r'\[Content_Types\]\.xml', '', text)
    text = re.sub(r'docProps/[\w.]+', '', text)
    text = re.sub(r'stylesWithEffects\.xml', '', text)
    # Remove xmlns attributes
    text = re.sub(r'xmlns:[\w]+=["\'][^"\']*["\']', '', text)
    text = re.sub(r'\bxmlns=["\'][^"\']*["\']', '', text)
    # Remove PK zip header artifacts from DOCX binary
    text = re.sub(r'PK[\x03\x04\x05\x06].*?(?=\n|$)', '', text)
    # ENGLISH ONLY: Remove ALL non-ASCII characters
    text = re.sub(r'[^\x20-\x7E\n\r\t]', ' ', text)
    # Collapse excessive whitespace
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
    # Remove lines that are mostly non-alphabetic (binary residue)
    lines = text.split('\n')
    clean_lines = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            clean_lines.append('')
            continue
        alpha_count = sum(1 for c in stripped if c.isalpha())
        if len(stripped) > 0 and alpha_count / len(stripped) > 0.3:
            # Additional gibberish check: reject lines with no real English words
            if not _is_line_gibberish(stripped):
                clean_lines.append(stripped)
    text = '\n'.join(clean_lines).strip()
    return text


def _is_line_gibberish(line):
    """Check if a line is truly random gibberish (not just poor OCR).
    Only rejects lines that are clearly machine-generated noise."""
    if not line or len(line.strip()) < 10:
        return False  # too short to judge — keep it
    words = re.findall(r'\b[a-zA-Z]{2,}\b', line)
    if not words:
        return len(line.strip()) > 30  # long line with zero words = likely binary
    # Check for extreme vowel absence — only flag if NO words have vowels
    vowels = set('aeiouAEIOU')
    no_vowel_count = 0
    for w in words:
        if len(w) >= 4 and not any(c in vowels for c in w):
            no_vowel_count += 1
    # Only gibberish if majority of 4+ letter words have zero vowels
    long_words = [w for w in words if len(w) >= 4]
    if long_words and no_vowel_count / len(long_words) > 0.6:
        return True
    return False


def enforce_english(text):
    """Final English enforcement pass — strips ALL non-ASCII characters,
    XML metadata, DOCX package paths, and machine-generated noise.
    Ensures only clean professional English text is returned."""
    if not text:
        return ""
    # Strip all non-ASCII characters (foreign scripts, corrupted unicode, symbols)
    text = re.sub(r'[^\x20-\x7E\n\r\t]', ' ', text)
    # Remove any remaining XML/HTML tags
    text = re.sub(r'<[^>]+>', ' ', text)
    # Remove DOCX internal paths that might have leaked
    text = re.sub(r'word/[\w._/-]+\.xml', '', text, flags=re.IGNORECASE)
    text = re.sub(r'_rels/[\w.]+', '', text)
    text = re.sub(r'\[Content_Types\]', '', text)
    text = re.sub(r'docProps/[\w.]+', '', text)
    text = re.sub(r'stylesWithEffects', '', text)
    # Remove encoded/binary junk
    text = re.sub(r'PK[\x03\x04\x05\x06].*?(?=\n|$)', '', text)
    # Remove xmlns artifacts
    text = re.sub(r'xmlns:[\w]+=["\'][^"\']*["\']', '', text)
    # Collapse whitespace
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
    return text.strip()


def is_text_readable(text):
    """Check if extracted text is actually readable English (not garbled encoding).
    Uses strict thresholds to catch corrupted font-mapped ASCII garbage."""
    if not text or len(text.strip()) < 20:
        return False
    # Only count 3+ letter words to avoid false positives from 'a', 'I', etc.
    common_words = {
        'the', 'and', 'for', 'that', 'was', 'are', 'with', 'his', 'they',
        'this', 'have', 'from', 'one', 'had', 'but', 'not', 'what', 'all',
        'were', 'when', 'your', 'can', 'said', 'there', 'each', 'which',
        'shall', 'party', 'agreement', 'between', 'herein', 'section', 'under',
        'any', 'such', 'will', 'may', 'has', 'been', 'other', 'than', 'its',
        'upon', 'more', 'into', 'made', 'after', 'date', 'terms', 'between',
        'property', 'document', 'contract', 'law', 'court', 'legal', 'notice',
        'person', 'company', 'amount', 'payment', 'period', 'right', 'clause',
        'tenant', 'landlord', 'rent', 'lease', 'agreement', 'signed', 'witness',
        'whereas', 'hereby', 'therefore', 'provided', 'subject', 'including',
    }
    words = re.findall(r'\b[a-zA-Z]{3,}\b', text.lower())
    if len(words) < 5:
        return False
    common_count = sum(1 for w in words if w in common_words)
    ratio = common_count / len(words)
    # Strict: at least 12% of 3+ letter words must be common English
    return ratio > 0.12


def try_ocr_extraction(file_obj):
    """Attempt OCR extraction from a PDF using pytesseract + pdf2image."""
    try:
        import pytesseract
        from pdf2image import convert_from_bytes
        file_obj.seek(0)
        pdf_bytes = file_obj.read()
        images = convert_from_bytes(pdf_bytes, dpi=300)
        ocr_text = ""
        for i, img in enumerate(images):
            page_text = pytesseract.image_to_string(img, lang='eng')
            if page_text.strip():
                ocr_text += page_text + "\n"
        return ocr_text.strip()
    except ImportError:
        logger.warning("OCR libraries (pytesseract/pdf2image) not fully available.")
        return ""
    except Exception as e:
        logger.warning(f"OCR extraction failed: {e}")
        return ""


def extract_text_from_file(file_obj):
    """Extract clean readable text from PDF, DOCX, or TXT files.
    Includes OCR fallback for scanned/image-based PDFs."""
    filename = file_obj.name.lower()
    text = ""

    # Try PDF first
    if filename.endswith('.pdf'):
        try:
            file_obj.seek(0)
            reader = PdfReader(file_obj)
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
        except Exception:
            pass

        # Check if PyPDF2 output is readable; if not, try OCR
        sanitized = sanitize_text(text)
        if not is_text_readable(sanitized):
            logger.info("PyPDF2 text not readable, attempting OCR extraction...")
            ocr_text = try_ocr_extraction(file_obj)
            if ocr_text and is_text_readable(ocr_text):
                text = ocr_text
                logger.info("OCR extraction successful.")
            else:
                # Return a clear English message instead of garbled text
                text = (
                    "This document appears to be a scanned image or uses custom font "
                    "encoding that could not be decoded. The extracted text was not "
                    "readable.\n\n"
                    "To analyze this document properly, please:\n"
                    "1. Install Tesseract OCR (https://github.com/tesseract-ocr/tesseract) "
                    "and ensure it is in your system PATH.\n"
                    "2. Install Poppler for PDF-to-image conversion "
                    "(https://github.com/oschwartz10612/poppler-windows/releases).\n"
                    "3. Re-upload the document for OCR-based text extraction.\n\n"
                    "Alternatively, you can convert the scanned PDF to a text-based PDF "
                    "using an online OCR tool and re-upload it."
                )
                return text

    # Try DOCX
    elif filename.endswith('.docx'):
        try:
            file_obj.seek(0)
            from docx import Document as DocxDocument
            doc = DocxDocument(file_obj)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = '\n'.join(paragraphs)
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text += '\n' + row_text
        except ImportError:
            # python-docx not installed, try raw extraction
            file_obj.seek(0)
            raw = file_obj.read().decode('utf-8', errors='ignore')
            text = sanitize_text(raw)
        except Exception:
            file_obj.seek(0)
            raw = file_obj.read().decode('utf-8', errors='ignore')
            text = sanitize_text(raw)

    # Try plain text
    elif filename.endswith('.txt'):
        try:
            file_obj.seek(0)
            text = file_obj.read().decode('utf-8', errors='ignore')
        except Exception:
            pass

    # Fallback: try PDF, then raw
    else:
        try:
            file_obj.seek(0)
            reader = PdfReader(file_obj)
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
        except Exception:
            file_obj.seek(0)
            text = file_obj.read().decode('utf-8', errors='ignore')

    # Always sanitize
    text = sanitize_text(text)

    # NOTE: Never replace extracted text with a static fallback message.
    # Even poor quality text contains patterns that the NLP engine can analyze
    # for document type, risk patterns, compliance checks, etc.
    # The quality notice is handled separately in the reconstruction engine.

    if not text.strip():
        text = "Document could not be parsed. The file may be a scanned image requiring OCR, or an unsupported format."

    return text


class DocumentListView(generics.ListAPIView):
    queryset = Document.objects.all().order_by('-uploaded_at')
    serializer_class = DocumentSerializer


class DocumentHistoryView(generics.ListAPIView):
    serializer_class = HistorySerializer
    def get_queryset(self):
        return History.objects.filter(document_id=self.kwargs['pk'])


class DocumentProcessView(APIView):
    permission_classes = [AllowAny]
    """Upload and fully analyze a legal document with AI reconstruction."""
    def post(self, request, *args, **kwargs):
        file_obj = request.FILES.get('file')
        if not file_obj:
            return Response({"error": "No file provided"}, status=status.HTTP_400_BAD_REQUEST)

        text_content = extract_text_from_file(file_obj)

        # AI Text Reconstruction — always clean the text
        reconstruction = nlp_engine.reconstruct_text(text_content)
        clean_text = reconstruction["reconstructed_text"]

        # Use reconstructed text for analysis if reconstruction was applied
        analysis_text = clean_text if reconstruction["reconstruction_applied"] else text_content

        # Full NLP analysis
        doc_type = nlp_engine.classify_document(text_content)
        summary = nlp_engine.generate_summary(analysis_text)
        result = predict_text(summary)
        risk_data = nlp_engine.analyze_risk(analysis_text)
        fraud_data = nlp_engine.detect_fraud(analysis_text)
        compliance_data = nlp_engine.check_compliance(analysis_text)
        simplified = nlp_engine.simplify_text(analysis_text)
        key_data = nlp_engine.extract_key_data(text_content)

        # AI Document Summary (structured)
        ai_summary = nlp_engine.generate_ai_document_summary(text_content, doc_type)

        document = Document.objects.create(
            owner=request.user if request.user.is_authenticated else None,
            original_filename=file_obj.name,
            content=text_content,
            summary=summary,
            simplified_text=simplified["simplified_text"],
            doc_type=doc_type,
            risk_score=risk_data["risk_score"],
            fraud_score=fraud_data["fraud_score"],
            compliance_score=compliance_data["score"],
            risk_analysis_json=risk_data,
            fraud_analysis_json=fraud_data,
            compliance_json=compliance_data,
            key_data_json=key_data,
        )
        History.objects.create(document=document, action="Document Uploaded & Analyzed")

        return Response({
            "document_id": document.id,
            "filename": file_obj.name,
            "content": enforce_english(clean_text[:5000]),
            "raw_content": enforce_english(text_content[:2000]),
            "summary": enforce_english(summary),
            "ml_score": result,
            "ml_result": result,
            "doc_type": doc_type,
            "risk": risk_data,
            "fraud": fraud_data,
            "compliance": compliance_data,
            "simplified": simplified,
            "key_data": key_data,
            "reconstruction": {
                "reconstructed_text": enforce_english(clean_text[:5000]),
                "applied": reconstruction["reconstruction_applied"],
                "method": reconstruction["method"],
                "quality": reconstruction["quality"],
            },
            "ai_summary": ai_summary,
        }, status=status.HTTP_200_OK)


class ReconstructView(APIView):
    """Force clean English reconstruction of document text."""
    def post(self, request, *args, **kwargs):
        doc_id = request.data.get('document_id')
        text = request.data.get('text', '')
        if doc_id and not text:
            try:
                doc = Document.objects.get(pk=doc_id)
                text = doc.content
            except Document.DoesNotExist:
                return Response({"error": "Document not found"}, status=status.HTTP_404_NOT_FOUND)
        if not text:
            return Response({"error": "No text provided"}, status=status.HTTP_400_BAD_REQUEST)

        result = nlp_engine.reconstruct_text(text)
        ai_summary = nlp_engine.generate_ai_document_summary(text)
        result["ai_summary"] = ai_summary
        return Response(result)


class DocumentChatView(APIView):
    """RAG-style Q&A against uploaded document."""
    def post(self, request, *args, **kwargs):
        doc_id = request.data.get('document_id')
        question = request.data.get('question', '')
        if not question:
            return Response({"error": "No question provided"}, status=status.HTTP_400_BAD_REQUEST)
        text = ""
        if doc_id:
            try:
                doc = Document.objects.get(pk=doc_id)
                text = doc.content
            except Document.DoesNotExist:
                pass
        result = nlp_engine.answer_question(text, question)
        # Enforce English on the answer
        if 'answer' in result:
            result['answer'] = enforce_english(result['answer'])
        return Response(result, status=status.HTTP_200_OK)


class RiskAnalysisView(APIView):
    """Clause-by-clause risk analysis."""
    def get(self, request, pk, *args, **kwargs):
        try:
            doc = Document.objects.get(pk=pk)
        except Document.DoesNotExist:
            return Response({"error": "Document not found"}, status=status.HTTP_404_NOT_FOUND)
        if doc.risk_analysis_json:
            return Response(doc.risk_analysis_json)
        result = nlp_engine.analyze_risk(doc.content)
        doc.risk_analysis_json = result
        doc.risk_score = result["risk_score"]
        doc.save()
        return Response(result)


class FraudDetectionView(APIView):
    """Fraud and missing field detection."""
    def get(self, request, pk, *args, **kwargs):
        try:
            doc = Document.objects.get(pk=pk)
        except Document.DoesNotExist:
            return Response({"error": "Document not found"}, status=status.HTTP_404_NOT_FOUND)
        if doc.fraud_analysis_json:
            return Response(doc.fraud_analysis_json)
        result = nlp_engine.detect_fraud(doc.content)
        doc.fraud_analysis_json = result
        doc.fraud_score = result["fraud_score"]
        doc.save()
        return Response(result)


class ComplianceCheckView(APIView):
    """Compliance checklist."""
    def get(self, request, pk, *args, **kwargs):
        try:
            doc = Document.objects.get(pk=pk)
        except Document.DoesNotExist:
            return Response({"error": "Document not found"}, status=status.HTTP_404_NOT_FOUND)
        if doc.compliance_json:
            return Response(doc.compliance_json)
        result = nlp_engine.check_compliance(doc.content)
        doc.compliance_json = result
        doc.compliance_score = result["score"]
        doc.save()
        return Response(result)


class SimplifyView(APIView):
    """Simplify legal text."""
    def post(self, request, *args, **kwargs):
        text = request.data.get('text', '')
        doc_id = request.data.get('document_id')
        if doc_id and not text:
            try:
                doc = Document.objects.get(pk=doc_id)
                text = doc.content
            except Document.DoesNotExist:
                pass
        if not text:
            return Response({"error": "No text provided"}, status=status.HTTP_400_BAD_REQUEST)
        result = nlp_engine.simplify_text(text)
        return Response(result)


class ClauseGeneratorView(APIView):
    """Generate legal clauses from templates."""
    def get(self, request, *args, **kwargs):
        templates = []
        for key, val in legal_templates.TEMPLATES.items():
            fields = legal_templates.get_template_fields(key)
            templates.append({"id": key, "title": val["title"], "fields": fields})
        return Response({"templates": templates})

    def post(self, request, *args, **kwargs):
        template_type = request.data.get('template_type')
        fields = request.data.get('fields', {})
        if not template_type:
            return Response({"error": "template_type required"}, status=status.HTTP_400_BAD_REQUEST)
        result = legal_templates.generate_clause(template_type, fields)
        if "error" in result:
            return Response(result, status=status.HTTP_400_BAD_REQUEST)
        if request.user.is_authenticated:
            GeneratedClause.objects.create(
                owner=request.user, template_type=template_type,
                title=result["title"], content=result["content"]
            )
        return Response(result)


class DocumentCompareView(APIView):
    """Compare two documents."""
    def post(self, request, *args, **kwargs):
        file1 = request.FILES.get('file1')
        file2 = request.FILES.get('file2')
        if not file1 or not file2:
            return Response({"error": "Two files required"}, status=status.HTTP_400_BAD_REQUEST)

        text1, text2 = extract_text_from_file(file1), extract_text_from_file(file2)
        result = nlp_engine.compare_documents(text1, text2)
        result["file1_name"] = file1.name
        result["file2_name"] = file2.name
        return Response(result)


class LawyerRecommendView(APIView):
    """Recommend lawyers based on document type and location."""
    def post(self, request, *args, **kwargs):
        doc_id = request.data.get('document_id')
        user_lat = request.data.get('latitude')
        user_lng = request.data.get('longitude')
        specialty = request.data.get('specialty', 'Civil Law')

        if doc_id:
            try:
                doc = Document.objects.get(pk=doc_id)
                doc_type = doc.doc_type or nlp_engine.classify_document(doc.content)
                specialty = nlp_engine.LAWYER_SPECIALTIES.get(doc_type, 'Civil Law')
            except Document.DoesNotExist:
                pass

        lawyers = lawyer_data.get_recommended_lawyers(
            specialty,
            float(user_lat) if user_lat else None,
            float(user_lng) if user_lng else None
        )
        return Response({"specialty": specialty, "lawyers": lawyers})


class TranslatorView(APIView):
    """Smart bilingual translator — auto-detects English or Kannada and translates."""
    def post(self, request, *args, **kwargs):
        text = request.data.get('text', '')
        doc_id = request.data.get('document_id')

        # Get text from document if doc_id provided
        if doc_id and not text:
            try:
                doc = Document.objects.get(pk=doc_id)
                text = doc.content
            except Document.DoesNotExist:
                pass

        if not text:
            return Response({"error": "No text provided for translation."}, status=status.HTTP_400_BAD_REQUEST)

        # Auto-detect language
        detected_lang = self._detect_language(text)
        target_lang = 'en' if detected_lang == 'kn' else 'kn'
        source_label = 'Kannada' if detected_lang == 'kn' else 'English'
        target_label = 'English' if target_lang == 'en' else 'Kannada'

        # Clean text before translation
        clean_text = sanitize_text(text)[:3000]

        # Attempt translation
        translated = self._translate(clean_text, detected_lang, target_lang)

        return Response({
            "translated_text": translated,
            "source_language": source_label,
            "target_language": target_label,
            "detected": source_label,
            "label": f"AI Smart Translation Output ({source_label} → {target_label})",
        })

    def _detect_language(self, text):
        """Detect if text is predominantly Kannada or English."""
        if not text:
            return 'en'
        # Count Kannada Unicode characters (U+0C80 to U+0CFF)
        kannada_chars = sum(1 for c in text if '\u0C80' <= c <= '\u0CFF')
        ascii_alpha = sum(1 for c in text if c.isascii() and c.isalpha())
        total = kannada_chars + ascii_alpha
        if total == 0:
            return 'en'
        if kannada_chars / total > 0.3:
            return 'kn'
        return 'en'

    def _translate(self, text, source, target):
        """Translate text using deep_translator (Google Translate) with fallback."""
        if not text.strip():
            return "No readable text available for translation."
        try:
            from deep_translator import GoogleTranslator
            # Translate in chunks if text is long
            chunks = [text[i:i+4500] for i in range(0, len(text), 4500)]
            translated_chunks = []
            for chunk in chunks:
                result = GoogleTranslator(source=source, target=target).translate(chunk)
                if result:
                    translated_chunks.append(result)
            translated = ' '.join(translated_chunks)
            if translated and translated.strip():
                return translated
            return "Translation could not be completed. Please try with different text."
        except ImportError:
            logger.warning("deep_translator not installed. Using fallback translation.")
            return self._fallback_translate(text, source, target)
        except Exception as e:
            logger.warning(f"Translation failed: {e}")
            return self._fallback_translate(text, source, target)

    def _fallback_translate(self, text, source, target):
        """Fallback when deep_translator is not available."""
        if target == 'en':
            return (
                "Translation service temporarily unavailable. "
                "The uploaded text appears to be in Kannada. "
                "Please install the 'deep-translator' package for full translation support: "
                "pip install deep-translator"
            )
        else:
            return (
                "Translation service temporarily unavailable. "
                "The uploaded text is in English. "
                "Please install the 'deep-translator' package for Kannada translation: "
                "pip install deep-translator"
            )


class SaveSignatureView(APIView):
    def post(self, request, *args, **kwargs):
        doc_id = request.data.get('document_id')
        sig = request.data.get('signature')
        if not doc_id or not sig:
            return Response({"error": "Missing document_id or signature"}, status=status.HTTP_400_BAD_REQUEST)
        try:
            doc = Document.objects.get(pk=doc_id)
            doc.signature = sig
            doc.save()
            History.objects.create(document=doc, action="Document Signed")
            return Response({"message": "Signature saved!"})
        except Document.DoesNotExist:
            return Response({"error": "Document not found"}, status=status.HTTP_404_NOT_FOUND)


class DownloadSignedView(APIView):
    def get(self, request, pk, *args, **kwargs):
        try:
            doc = Document.objects.get(pk=pk)
        except Document.DoesNotExist:
            return HttpResponse("Not found.", status=404)
        if not doc.signature:
            return HttpResponse("Not signed.", status=400)
        buf = io.BytesIO()
        p = canvas.Canvas(buf, pagesize=letter)
        w, h = letter
        p.setFont("Helvetica-Bold", 16)
        p.drawString(40, h-40, "Signed Document")
        p.setFont("Helvetica", 10)
        p.drawString(40, h-55, f"File: {doc.original_filename}")
        p.line(40, h-65, w-40, h-65)
        t = p.beginText(40, h-90)
        t.setFont("Helvetica", 10)
        for line in doc.content.split('\n'):
            for wr in textwrap.wrap(line, width=100):
                t.textLine(wr)
        p.drawText(t)
        try:
            sig_data = doc.signature.split(',')[1]
            sig_img = base64.b64decode(sig_data)
            p.drawString(40, 100, "Signed:")
            p.drawImage(ImageReader(io.BytesIO(sig_img)), x=40, y=110, width=150, height=75, mask='auto')
        except Exception:
            pass
        p.showPage()
        p.save()
        buf.seek(0)
        resp = HttpResponse(buf, content_type='application/pdf')
        resp['Content-Disposition'] = f'attachment; filename="signed_{doc.original_filename}"'
        return resp


class DownloadSummaryPDFView(APIView):
    def get(self, request, pk, *args, **kwargs):
        try:
            doc = Document.objects.get(pk=pk)
        except Document.DoesNotExist:
            return HttpResponse("Not found.", status=404)
        buf = io.BytesIO()
        p = canvas.Canvas(buf, pagesize=letter)
        w, h = letter
        p.setFont("Helvetica-Bold", 16)
        p.drawString(40, h-40, "Legal Document Analysis Report")
        p.setFont("Helvetica", 10)
        p.drawString(40, h-55, f"File: {doc.original_filename}")
        p.drawString(40, h-68, f"Type: {doc.doc_type or 'N/A'} | Risk: {doc.risk_score}% | Compliance: {doc.compliance_score}%")
        p.line(40, h-78, w-40, h-78)
        t = p.beginText(40, h-100)
        t.setFont("Helvetica-Bold", 12)
        t.textLine("SUMMARY:")
        t.setFont("Helvetica", 10)
        t.textLine("")
        for line in (doc.summary or "N/A").split('\n'):
            for wr in textwrap.wrap(line, width=100):
                t.textLine(wr)
        t.textLine("")
        t.setFont("Helvetica-Bold", 12)
        t.textLine("SIMPLIFIED TEXT:")
        t.setFont("Helvetica", 10)
        t.textLine("")
        for line in (doc.simplified_text or "N/A")[:2000].split('\n'):
            for wr in textwrap.wrap(line, width=100):
                t.textLine(wr)
        p.drawText(t)
        p.showPage()
        p.save()
        buf.seek(0)
        resp = HttpResponse(buf, content_type='application/pdf')
        resp['Content-Disposition'] = f'attachment; filename="report_{doc.original_filename}.pdf"'
        return resp


class DownloadTranslatedPDFView(APIView):
    def post(self, request, *args, **kwargs):
        text = request.data.get('text', '')
        if not text:
            return HttpResponse("No text.", status=400)
        buf = io.BytesIO()
        p = canvas.Canvas(buf, pagesize=letter)
        w, h = letter
        p.setFont("Helvetica-Bold", 16)
        p.drawString(40, h-40, "Translated Document")
        p.line(40, h-55, w-40, h-55)
        t = p.beginText(40, h-80)
        try:
            t.setFont("NotoSansKannada", 10)
        except Exception:
            t.setFont("Helvetica", 10)
        for line in text.split('\n'):
            for wr in textwrap.wrap(line, width=80):
                t.textLine(wr)
        p.drawText(t)
        p.showPage()
        p.save()
        buf.seek(0)
        resp = HttpResponse(buf, content_type='application/pdf')
        resp['Content-Disposition'] = 'attachment; filename="translated.pdf"'
        return resp